from __future__ import annotations

import re
from collections import Counter
from pathlib import Path

from .schemas import Document

SUPPORTED = {".pdf", ".txt", ".md", ".markdown", ".docx"}


def load_documents(paths: list[str | Path]) -> list[Document]:
    """Load every supported file under the given files/directories."""
    files: list[Path] = []
    for raw in paths:
        p = Path(raw)
        if p.is_dir():
            files.extend(sorted(f for f in p.rglob("*") if f.suffix.lower() in SUPPORTED))
        elif p.is_file() and p.suffix.lower() in SUPPORTED:
            files.append(p)
    docs: list[Document] = []
    for f in files:
        try:
            docs.append(_load_one(f))
        except Exception as exc:  # keep going; one bad file shouldn't abort
            docs.append(Document(name=f.name, text=f"[parse error: {exc}]"))
    return [d for d in docs if d.text and not d.text.startswith("[parse error")]


def _load_one(path: Path) -> Document:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text, n_pages, sections = _parse_pdf(path)
    elif suffix in {".md", ".markdown"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
        n_pages, sections = 0, _headings_from_markdown(text)
    elif suffix == ".docx":
        text, sections = _parse_docx(path)
        n_pages = 0
    else:  # .txt and anything else text-like
        text = path.read_text(encoding="utf-8", errors="ignore")
        n_pages, sections = 0, _headings_from_plaintext(text)
    return Document(name=path.name, text=text, n_pages=n_pages, sections=sections)


# --------------------------------------------------------------------------- #
# PDF (pymupdf, lazy-imported)
# --------------------------------------------------------------------------- #
def _parse_pdf(path: Path) -> tuple[str, int, list[str]]:
    try:
        import fitz  # pymupdf
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("Install `pymupdf` to parse PDF files.") from exc

    doc = fitz.open(path)
    page_texts: list[str] = []
    lines: list[tuple[str, float, int]] = []  # (text, font_size, flags)
    for page in doc:
        page_texts.append(page.get_text("text"))
        data = page.get_text("dict")
        for block in data.get("blocks", []):
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                line_text = "".join(s.get("text", "") for s in spans).strip()
                if not line_text:
                    continue
                size = max((s.get("size", 0.0) for s in spans), default=0.0)
                flags = max((s.get("flags", 0) for s in spans), default=0)
                lines.append((line_text, round(size, 1), flags))
    n_pages = doc.page_count
    doc.close()
    return "\n".join(page_texts), n_pages, _headings_from_spans(lines)


def _headings_from_spans(lines: list[tuple[str, float, int]]) -> list[str]:
    if not lines:
        return []
    body_size = Counter(size for _, size, _ in lines).most_common(1)[0][0]
    headings: list[str] = []
    for text, size, flags in lines:
        if not _looks_like_heading(text):
            continue
        words = text.split()
        is_bigger = size >= body_size * 1.15
        is_bold = bool(flags & 16)  # pymupdf bold bit
        is_numbered = bool(re.match(r"^\d+(\.\d+)*\s+\S", text))
        if is_bigger or (is_bold and len(words) <= 10) or (is_numbered and is_bigger):
            headings.append(text)
    return _dedupe(headings)


def _looks_like_heading(text: str) -> bool:
    """Quality gate for candidate headings (PDF/plaintext heuristics).

    Rejects the noise that wrecks section detection on real PDFs: page numbers,
    bare years ("2023"), percentages ("100%"), figure/table labels, and other
    digit- or symbol-dominated fragments. A real heading is mostly letters and
    contains at least one genuine word.
    """
    t = text.strip()
    if not (3 <= len(t) <= 120):
        return False
    words = t.split()
    if not (1 <= len(words) <= 14):
        return False
    if t.endswith((".", ",", ";", ":")):
        return False
    alpha = sum(c.isalpha() for c in t)
    digits = sum(c.isdigit() for c in t)
    if alpha <= digits or alpha / len(t) < 0.5:
        return False
    # require at least one real word (>= 3 letters), e.g. not "A1 B2 C3"
    if not any(sum(c.isalpha() for c in w) >= 3 for w in words):
        return False
    return True


# --------------------------------------------------------------------------- #
# docx (optional dependency)
# --------------------------------------------------------------------------- #
def _parse_docx(path: Path) -> tuple[str, list[str]]:
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise RuntimeError("Install `python-docx` (extra: docx) to parse .docx files.") from exc
    document = docx.Document(str(path))
    parts, sections = [], []
    for para in document.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        parts.append(t)
        if (para.style and para.style.name and para.style.name.lower().startswith("heading")):
            sections.append(t)
    return "\n".join(parts), _dedupe(sections)


# --------------------------------------------------------------------------- #
# Plain-text / markdown heading heuristics
# --------------------------------------------------------------------------- #
def _headings_from_markdown(text: str) -> list[str]:
    out = [m.group(1).strip() for m in re.finditer(r"^#{1,6}\s+(.+)$", text, re.MULTILINE)]
    return _dedupe(out)


def _headings_from_plaintext(text: str) -> list[str]:
    out: list[str] = []
    for line in text.splitlines():
        s = line.strip()
        if not _looks_like_heading(s) or len(s) > 90:
            continue
        if re.match(r"^\d+(\.\d+)*\s+\S", s) or (s.isupper() and len(s) > 3):
            out.append(s)
    return _dedupe(out)


def _dedupe(items: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for it in items:
        key = it.lower()
        if key not in seen:
            seen.add(key)
            out.append(it)
    return out
